from dataclasses import dataclass
from operator import index
from tkinter import font
import pandas as pd
import numpy as np
import os
import math
import scipy
import matplotlib.pyplot as plt
import re
import seaborn as sns
import statsmodels.api as sm
import scipy.stats as stats
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix
from sklearn.metrics import classification_report
from sklearn.metrics import accuracy_score
from sklearn import metrics
from sklearn import preprocessing
from sklearn.model_selection import cross_val_score
from docx import Document
from docx2pdf import convert
from docx.shared import Inches
from scipy.stats import shapiro
from scipy.stats import kstest
from scipy.stats import shapiro
from scipy.stats import kstest
from scipy.stats import normaltest
from scipy.stats import anderson
from scipy.stats import ttest_ind
import scipy.stats as stats
from scipy.stats import spearmanr
from scipy.stats import pearsonr
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix
from sklearn.metrics import classification_report
from sklearn.metrics import accuracy_score
from sklearn import metrics
from sklearn import preprocessing
from sklearn.model_selection import cross_val_score
from sklearn.impute import SimpleImputer, MissingIndicator
from sklearn.model_selection import train_test_split
from sklearn.pipeline import FeatureUnion, make_pipeline
from sklearn import tree 
from sklearn.tree import export_graphviz
from io import StringIO  
from IPython.display import Image  
import pydotplus
from sklearn.model_selection import cross_validate

from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
import shap
from sklearn.svm import SVC
from sklearn import svm
from sklearn.svm import LinearSVC
from sklearn.neighbors import KNeighborsClassifier as KNN
from sklearn.ensemble import RandomForestClassifier as RF
from sklearn.ensemble import GradientBoostingClassifier as GBC
from sklearn.ensemble import AdaBoostClassifier as ABC
from sklearn.ensemble import BaggingClassifier as BGC
from sklearn import decomposition, datasets
from sklearn.model_selection import GridSearchCV
from sklearn.pipeline import Pipeline
from sklearn.model_selection import RandomizedSearchCV
import joblib
from sklearn.metrics import make_scorer
from sklearn.metrics import recall_score
from sklearn.model_selection import cross_validate

class classifier():
    def __init__(self, dataframe_1, path):
        self.dataframe = dataframe_1
        
        self.path = path

    def feature_transform(self):
        data= self.dataframe
        

        path = self.path

        new_dict = {'ID':[],'Subgroup':[], 'SF_S1_E2':[], 'SF_S2_E2':[], 'SF_S3_E2':[], 'SF_S1_P4':[], 'SF_S2_P4':[], 'SF_S3_P4':[], 'SC_S1_E2':[], 'SC_S2_E2':[], 'SC_S3_E2':[], 'SC_S1_P4':[], 'SC_S2_P4':[], 'SC_S3_P4':[], 
'BC_S1_E2':[], 'BC_S2_E2':[], 'BC_S3_E2':[], 'BC_S1_P4':[], 'BC_S2_P4':[], 'BC_S3_P4':[], 'SF_S1-3_E2_RelativeChange':[], 'SC_S1-3_E2_RelativeChange':[],'BC_S1-3_E2_RelativeChange':[], 'SF_S2-3_E2_RelativeChange':[], 
'SC_S2-3_E2_RelativeChange':[], 'BC_S2-3_E2_RelativeChange':[], 'SF_S1-2_E2_RelativeChange':[], 'SC_S1-2_E2_RelativeChange':[], 'BC_S1-2_E2_RelativeChange':[], 'SF_S1-3_E2_AbsoluteChange':[], 'SC_S1-3_E2_AbsoluteChange':[], 'BC_S1-3_E2_AbsoluteChange':[], 
'SF_S2-3_E2_AbsoluteChange':[], 'SC_S2-3_E2_AbsoluteChange':[], 'BC_S2-3_E2_AbsoluteChange':[], 'SF_S1-2_E2_AbsoluteChange':[], 'SC_S1-2_E2_AbsoluteChange':[], 'BC_S1-2_E2_AbsoluteChange':[], 
'SF_S1-3_P4_RelativeChange':[], 'SC_S1-3_P4_RelativeChange':[], 'BC_S1-3_P4_RelativeChange':[], 'SF_S2-3_P4_RelativeChange':[], 'SC_S2-3_P4_RelativeChange':[], 'BC_S2-3_P4_RelativeChange':[], 
'SF_S1-2_P4_RelativeChange':[], 'SC_S1-2_P4_RelativeChange':[], 'BC_S1-2_P4_RelativeChange':[], 'SF_S1-3_P4_AbsoluteChange':[], 'SC_S1-3_P4_AbsoluteChange':[], 'BC_S1-3_P4_AbsoluteChange':[], 
'SF_S2-3_P4_AbsoluteChange':[], 'SC_S2-3_P4_AbsoluteChange':[], 'BC_S2-3_P4_AbsoluteChange':[], 'SF_S1-2_P4_AbsoluteChange':[], 'SC_S1-2_P4_AbsoluteChange':[], 'BC_S1-2_P4_AbsoluteChange':[],
'SF_S1_E2:P4_Ratio':[], 'SC_S1_E2:P4_Ratio':[], 'BC_S1_E2:P4_Ratio':[],
'SF_S2_E2:P4_Ratio':[], 'SC_S2_E2:P4_Ratio':[], 'BC_S2_E2:P4_Ratio':[], 'SF_S3_E2:P4_Ratio':[], 'SC_S3_E2:P4_Ratio':[], 'BC_S3_E2:P4_Ratio':[], 
'SF_E2:P4_Ratio_mean':[], 'SC_E2:P4_Ratio_mean':[], 'SF_E2:P4_Ratio_std':[], 'SC_E2:P4_Ratio_std':[], 'SF_S2_E2_PredictedChange':[], 
'SC_S2_E2_PredictedChange':[], 'SF_S2_P4_PredictedChange':[], 'SC_S2_P4_PredictedChange':[], 'SF_S3_E2_PredictedChange':[], 'SC_S3_E2_PredictedChange':[], 
'SF_S3_P4_PredictedChange':[], 'SC_S3_P4_PredictedChange':[],
'BC_E2_Decision':[], 'BC_P4_Decision':[]}

#'BC_E2_Decision_f':[], 'BC_P4_Decision_f':[], 'BC_E2_Decision_1000':[], 'BC_P4_Decision_1000':[]

        patients = data['ID'].unique()

        

        for p in patients:
            #print(p)
            new_dict['ID'].append(p)
            subgroup = data.loc[data['ID']==p,'Subgroup'].values[0]
            new_dict['Subgroup'].append(subgroup)
            
            sf_s1_e2 = data.loc[(data['ID']==p) & (data['Round']==1.0),'SF_Estradiol (pg/mL)'].values
            if len(sf_s1_e2)==0:
                sf_s1_e2=[np.nan]
            new_dict['SF_S1_E2'].append(sf_s1_e2[0])
            sf_s2_e2 = data.loc[(data['ID']==p) & (data['Round']==2.0),'SF_Estradiol (pg/mL)'].values
            if len(sf_s2_e2)==0:
                sf_s2_e2 = [np.nan]
            new_dict['SF_S2_E2'].append(sf_s2_e2[0])
            sf_s3_e2 = data.loc[(data['ID']==p) & (data['Round']==3.0),'SF_Estradiol (pg/mL)'].values
            if len(sf_s3_e2)==0:
                sf_s3_e2 = [np.nan]
            new_dict['SF_S3_E2'].append(sf_s3_e2[0])
            sf_s1_p4 = data.loc[(data['ID']==p) & (data['Round']==1.0),'SF_Progesterone (pg/mL)'].values
            if len(sf_s1_p4)==0:
                sf_s1_p4 = [np.nan]
            new_dict['SF_S1_P4'].append(sf_s1_p4[0])
            sf_s2_p4 = data.loc[(data['ID']==p) & (data['Round']==2.0),'SF_Progesterone (pg/mL)'].values
            if len(sf_s2_p4)==0:
                sf_s2_p4 = [np.nan]
            new_dict['SF_S2_P4'].append(sf_s2_p4[0])
            sf_s3_p4 = data.loc[(data['ID']==p) & (data['Round']==3.0),'SF_Progesterone (pg/mL)'].values
            if len(sf_s3_p4)==0:
                sf_s3_p4 = [np.nan]
            new_dict['SF_S3_P4'].append(sf_s3_p4[0])
            sc_s1_e2 = data.loc[(data['ID']==p) & (data['Round']==1.0),'SC_Estradiol (pg/mL)'].values
            if len(sc_s1_e2)==0:
                sc_s1_e2 = [np.nan]
            new_dict['SC_S1_E2'].append(sc_s1_e2[0])
            sc_s2_e2 = data.loc[(data['ID']==p) & (data['Round']==2.0),'SC_Estradiol (pg/mL)'].values
            if len(sc_s2_e2)==0:
                sc_s2_e2 = [np.nan]
            new_dict['SC_S2_E2'].append(sc_s2_e2[0])
            sc_s3_e2 = data.loc[(data['ID']==p) & (data['Round']==3.0),'SC_Estradiol (pg/mL)'].values
            if len(sc_s3_e2)==0:
                sc_s3_e2 = [np.nan]
            new_dict['SC_S3_E2'].append(sc_s3_e2[0])
            sc_s1_p4 = data.loc[(data['ID']==p) & (data['Round']==1.0),'SC_Progesterone (pg/mL)'].values
            if len(sc_s1_p4)==0:
                sc_s1_p4 = [np.nan]
            new_dict['SC_S1_P4'].append(sc_s1_p4[0])
            sc_s2_p4 = data.loc[(data['ID']==p) & (data['Round']==2.0),'SC_Progesterone (pg/mL)'].values
            if len(sc_s2_p4)==0:
                sc_s2_p4 = [np.nan]
            new_dict['SC_S2_P4'].append(sc_s2_p4[0])
            sc_s3_p4 = data.loc[(data['ID']==p) & (data['Round']==3.0),'SC_Progesterone (pg/mL)'].values
            if len(sc_s3_p4)==0:
                sc_s3_p4 = [np.nan]
            new_dict['SC_S3_P4'].append(sc_s3_p4[0])
            bc_s1_e2 = data.loc[(data['ID']==p) & (data['Round']==1.0),'BC_Estradiol (pg/mL)'].values
            if len(bc_s1_e2)==0:
                bc_s1_e2 = [np.nan]
            new_dict['BC_S1_E2'].append(bc_s1_e2[0])
            bc_s2_e2 = data.loc[(data['ID']==p) & (data['Round']==2.0),'BC_Estradiol (pg/mL)'].values
            if len(bc_s2_e2)==0:
                bc_s2_e2 = [np.nan]
            new_dict['BC_S2_E2'].append(bc_s2_e2[0])
            bc_s3_e2 = data.loc[(data['ID']==p) & (data['Round']==3.0),'BC_Estradiol (pg/mL)'].values
            if len(bc_s3_e2)==0:
                bc_s3_e2 = [np.nan]
            new_dict['BC_S3_E2'].append(bc_s3_e2[0])
            bc_s1_p4 = data.loc[(data['ID']==p) & (data['Round']==1.0),'BC_Progesterone (pg/mL)'].values
            if len(bc_s1_p4)==0:
                bc_s1_p4 = [np.nan]
            new_dict['BC_S1_P4'].append(bc_s1_p4[0])
            bc_s2_p4 = data.loc[(data['ID']==p) & (data['Round']==2.0),'BC_Progesterone (pg/mL)'].values
            if len(bc_s2_p4)==0:
                bc_s2_p4 = [np.nan]
            new_dict['BC_S2_P4'].append(bc_s2_p4[0])
            bc_s3_p4 = data.loc[(data['ID']==p) & (data['Round']==3.0),'BC_Progesterone (pg/mL)'].values
            if len(bc_s3_p4)==0:
                bc_s3_p4 = [np.nan]
            new_dict['BC_S3_P4'].append(bc_s3_p4[0])

            
            sf_s1_e2_ratio = sf_s1_e2[0]/sf_s1_p4[0]
            new_dict['SF_S1_E2:P4_Ratio'].append(sf_s1_e2_ratio)
            sf_s2_e2_ratio = sf_s2_e2[0]/sf_s2_p4[0]
            new_dict['SF_S2_E2:P4_Ratio'].append(sf_s2_e2_ratio)
            sf_s3_e2_ratio = sf_s3_e2[0]/sf_s3_p4[0]
            new_dict['SF_S3_E2:P4_Ratio'].append(sf_s3_e2_ratio)
            
            sc_s1_e2_ratio = sc_s1_e2[0]/sc_s1_p4[0]
            new_dict['SC_S1_E2:P4_Ratio'].append(sc_s1_e2_ratio)
            sc_s2_e2_ratio = sc_s2_e2[0]/sc_s2_p4[0]
            new_dict['SC_S2_E2:P4_Ratio'].append(sc_s2_e2_ratio)
            sc_s3_e2_ratio = sc_s3_e2[0]/sc_s3_p4[0]
            new_dict['SC_S3_E2:P4_Ratio'].append(sc_s3_e2_ratio)
            
            bc_s1_e2_ratio = bc_s1_e2[0]/bc_s1_p4[0]
            new_dict['BC_S1_E2:P4_Ratio'].append(bc_s1_e2_ratio)
            bc_s2_e2_ratio = bc_s2_e2[0]/bc_s2_p4[0]
            new_dict['BC_S2_E2:P4_Ratio'].append(bc_s2_e2_ratio)
            bc_s3_e2_ratio = bc_s3_e2[0]/bc_s3_p4[0]
            new_dict['BC_S3_E2:P4_Ratio'].append(bc_s3_e2_ratio)

            new_dict['SF_E2:P4_Ratio_mean'].append(np.nanmean([sf_s1_e2_ratio,sf_s2_e2_ratio,sf_s3_e2_ratio]))
            new_dict['SC_E2:P4_Ratio_mean'].append(np.nanmean([sc_s1_e2_ratio,sc_s2_e2_ratio,sc_s3_e2_ratio]))
            new_dict['SF_E2:P4_Ratio_std'].append(np.std([sf_s1_e2_ratio,sf_s2_e2_ratio,sf_s3_e2_ratio]))
            new_dict['SC_E2:P4_Ratio_std'].append(np.std([sc_s1_e2_ratio,sc_s2_e2_ratio,sc_s3_e2_ratio]))

          
            new_dict['SF_S1-3_E2_RelativeChange'].append((sf_s3_e2[0] - sf_s1_e2[0])/sf_s1_e2[0])
            new_dict['SC_S1-3_E2_RelativeChange'].append((sc_s3_e2[0] - sc_s1_e2[0])/sc_s1_e2[0])
            new_dict['BC_S1-3_E2_RelativeChange'].append((bc_s3_e2[0] - bc_s1_e2[0])/bc_s1_e2[0])
            new_dict['SF_S2-3_E2_RelativeChange'].append((sf_s3_e2[0] - sf_s2_e2[0])/sf_s2_e2[0])
            new_dict['SC_S2-3_E2_RelativeChange'].append((sc_s3_e2[0] - sc_s2_e2[0])/sc_s2_e2[0])
            new_dict['BC_S2-3_E2_RelativeChange'].append((bc_s3_e2[0] - bc_s2_e2[0])/bc_s2_e2[0])
            new_dict['SF_S1-3_P4_RelativeChange'].append((sf_s3_p4[0] - sf_s1_p4[0])/sf_s1_p4[0])
            new_dict['SC_S1-3_P4_RelativeChange'].append((sc_s3_p4[0] - sc_s1_p4[0])/sc_s1_p4[0])
            new_dict['BC_S1-3_P4_RelativeChange'].append((bc_s3_p4[0] - bc_s1_p4[0])/bc_s1_p4[0])
            new_dict['SF_S2-3_P4_RelativeChange'].append((sf_s3_p4[0] - sf_s2_p4[0])/sf_s2_p4[0])
            new_dict['SC_S2-3_P4_RelativeChange'].append((sc_s3_p4[0] - sc_s2_p4[0])/sc_s2_p4[0])
            new_dict['BC_S2-3_P4_RelativeChange'].append((bc_s3_p4[0] - bc_s2_p4[0])/bc_s2_p4[0])
            new_dict['SF_S1-2_E2_RelativeChange'].append((sf_s2_e2[0] - sf_s1_e2[0])/sf_s1_e2[0])
            new_dict['SC_S1-2_E2_RelativeChange'].append((sc_s2_e2[0] - sc_s1_e2[0])/sc_s1_e2[0])
            new_dict['BC_S1-2_E2_RelativeChange'].append((bc_s2_e2[0] - bc_s1_e2[0])/bc_s1_e2[0])
            new_dict['SF_S1-2_P4_RelativeChange'].append((sf_s2_p4[0] - sf_s1_p4[0])/sf_s1_p4[0])
            new_dict['SC_S1-2_P4_RelativeChange'].append((sc_s2_p4[0] - sc_s1_p4[0])/sc_s1_p4[0])
            new_dict['BC_S1-2_P4_RelativeChange'].append((bc_s2_p4[0] - bc_s1_p4[0])/bc_s1_p4[0])
            new_dict['SF_S1-3_E2_AbsoluteChange'].append(sf_s3_e2[0] - sf_s1_e2[0])
            new_dict['SC_S1-3_E2_AbsoluteChange'].append(sc_s3_e2[0] - sc_s1_e2[0])
            new_dict['BC_S1-3_E2_AbsoluteChange'].append(bc_s3_e2[0] - bc_s1_e2[0])
            new_dict['SF_S2-3_E2_AbsoluteChange'].append(sf_s3_e2[0] - sf_s2_e2[0])
            new_dict['SC_S2-3_E2_AbsoluteChange'].append(sc_s3_e2[0] - sc_s2_e2[0])
            new_dict['BC_S2-3_E2_AbsoluteChange'].append(bc_s3_e2[0] - bc_s2_e2[0])
            new_dict['SF_S1-3_P4_AbsoluteChange'].append(sf_s3_p4[0] - sf_s1_p4[0])
            new_dict['SC_S1-3_P4_AbsoluteChange'].append(sc_s3_p4[0] - sc_s1_p4[0])
            new_dict['BC_S1-3_P4_AbsoluteChange'].append(bc_s3_p4[0] - bc_s1_p4[0])
            new_dict['SF_S2-3_P4_AbsoluteChange'].append(sf_s3_p4[0] - sf_s2_p4[0])
            new_dict['SC_S2-3_P4_AbsoluteChange'].append(sc_s3_p4[0] - sc_s2_p4[0])
            new_dict['BC_S2-3_P4_AbsoluteChange'].append(bc_s3_p4[0] - bc_s2_p4[0])
            new_dict['SF_S1-2_E2_AbsoluteChange'].append(sf_s2_e2[0] - sf_s1_e2[0])
            new_dict['SC_S1-2_E2_AbsoluteChange'].append(sc_s2_e2[0] - sc_s1_e2[0])
            new_dict['BC_S1-2_E2_AbsoluteChange'].append(bc_s2_e2[0] - bc_s1_e2[0])
            new_dict['SF_S1-2_P4_AbsoluteChange'].append(sf_s2_p4[0] - sf_s1_p4[0])
            new_dict['SC_S1-2_P4_AbsoluteChange'].append(sc_s2_p4[0] - sc_s1_p4[0])
            new_dict['BC_S1-2_P4_AbsoluteChange'].append(bc_s2_p4[0] - bc_s1_p4[0])
            #print(p, data.loc[data.ID==p])
            
                
            '''new_dict['SF_S2_E2_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==2.0),'SF_Estradiol (pg/mL)_predictedchange'].values[0])
            new_dict['SC_S2_E2_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==2.0),'SC_Estradiol (pg/mL)_predictedchange'].values[0])
            new_dict['SF_S3_E2_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==3.0),'SF_Estradiol (pg/mL)_predictedchange'].values[0])
            new_dict['SC_S3_E2_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==3.0),'SC_Estradiol (pg/mL)_predictedchange'].values[0])
            new_dict['SF_S2_P4_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==2.0),'SF_Progesterone (pg/mL)_predictedchange'].values[0])
            new_dict['SC_S2_P4_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==2.0),'SC_Progesterone (pg/mL)_predictedchange'].values[0])
            new_dict['SF_S3_P4_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==3.0),'SF_Progesterone (pg/mL)_predictedchange'].values[0])
            new_dict['SC_S3_P4_PredictedChange'].append(data.loc[(data['ID']==p) & (data['Round']==3.0),'SC_Progesterone (pg/mL)_predictedchange'].values[0])'''
            
            # FILL NA BC TRIGGER
            
            

            data['BC_E_TRIGGER'].replace(np.nan, 0, inplace=True)
            data['3class_trigger'].replace(np.nan, 0, inplace=True)
            #print(p, data.loc[data.ID==p, 'BC_E_TRIGGER'])
            if len(data.loc[(data.ID==p) & (data.Round==3.0)])>0:
                bc_e2_decision = data.loc[(data['ID']==p) & (data.Round==3.0),'BC_E_TRIGGER'].values[0]
                bc_p4_decision = data.loc[(data['ID']==p) & (data.Round==3.0), '3class_trigger'].values[0]
                '''bc_e2_decision_f = data.loc[(data['ID']==p) & (data.Round==3.0),'BC_E_TRIGGER_f'].values[0]
                bc_p4_decision_f = data.loc[(data['ID']==p) & (data.Round==3.0), 'BC_TRIGGER_f'].values[0]
                bc_e2_decision_1000 = data.loc[(data['ID']==p) & (data.Round==3.0),'BC_E_TRIGGER_1000'].values[0]
                bc_p4_decision_1000 = data.loc[(data['ID']==p) & (data.Round==3.0), 'BC_TRIGGER_1000'].values[0]'''
            else:
                bc_e2_decision = 0
                bc_p4_decision = 0
                '''bc_e2_decision_f = 0
                bc_p4_decision_f = 0
                bc_e2_decision_1000 = 0
                bc_p4_decision_1000 = 0'''

            #print('p trigger values', data['BC_TRIGGER'].value_counts())
            '''elif len(data.loc[(data['ID']==p) & (data['Round'] ==2.0)])>0:
                bc_e2_decision = data.loc[(data['ID']==p) & (data['Round'] ==2.0), 'BC_E_TRIGGER'].values[0]
                bc_p4_decision = data.loc[(data['ID']==p) & (data['Round'] ==2.0), 'BC_TRIGGER'].values[0]
            else:
                bc_e2_decision = data.loc[(data['ID']==p) & (data['Round'] ==1.0), 'BC_E_TRIGGER'].values[0]
                bc_p4_decision = data.loc[(data['ID']==p) & (data['Round'] ==1.0), 'BC_TRIGGER'].values[0]'''
            #print(bc_e2_decision)
            new_dict['BC_E2_Decision'].append(bc_e2_decision) 
            new_dict['BC_P4_Decision'].append(bc_p4_decision)
            '''new_dict['BC_E2_Decision_f'].append(bc_e2_decision_f)
            new_dict['BC_P4_Decision_f'].append(bc_p4_decision_f)
            new_dict['BC_E2_Decision_1000'].append(bc_e2_decision_1000)
            new_dict['BC_P4_Decision_1000'].append(bc_p4_decision_1000)'''

        data2=data

        data = pd.DataFrame.from_dict(new_dict, orient='index')
        data = data.transpose()

        #data2.replace(to_replace=['0-25','25-50', '50-75', '75-100'], value=[1,2,3,4], inplace=True)

        data2.Round.astype(float)


        for p in list(data2['ID'].unique()):
            
            
            if len(data2.loc[(data2['ID']==p) & (data2.Round == 1.0)])>0:
                data.loc[(data['ID']==p), 'Saliva:Serum_S1_P4_Ratio']=data2.loc[(data2['ID']==p) & (data2.Round == 1.0), 'SC:BC_P4_Ratio'].values[0]
                data.loc[(data['ID']==p), 'Saliva:Serum_S1_E2_Ratio']=data2.loc[(data2['ID']==p) & (data2.Round == 1.0), 'SC:BC_E2_Ratio'].values[0]
                data.loc[(data['ID']==p), 'Saliva:Serum_S1_P4_Ratio_bucket']=data2.loc[(data2['ID']==p) & (data2.Round == 1.0), 'SC:BC_P4_Ratio_Bucket'].values[0]
                data.loc[(data['ID']==p), 'Saliva:Serum_S1_E2_Ratio_bucket']=data2.loc[(data2['ID']==p) & (data2.Round == 1.0), 'SC:BC_E2_Ratio_Bucket'].values[0]


            
            data.loc[(data['ID']==p), 'SF_E2_mean']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)'].astype(float).mean()
            data.loc[(data['ID']==p), 'SF_P4_mean']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)'].astype(float).mean()
            data.loc[(data['ID']==p), 'SC_E2_mean']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)'].astype(float).mean()
            data.loc[(data['ID']==p), 'SC_P4_mean']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)'].astype(float).mean()
            data.loc[(data['ID']==p), 'SF_E2_std']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)'].astype(float).std()
            data.loc[(data['ID']==p), 'SF_P4_std']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)'].astype(float).std()
            data.loc[(data['ID']==p), 'SC_E2_std']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)'].astype(float).std()
            data.loc[(data['ID']==p), 'SC_P4_std']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)'].astype(float).std()
            
            data.loc[(data['ID']==p), 'SF_E2_SE']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)'].astype(float).sem()
            data.loc[(data['ID']==p), 'SF_P4_SE']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)'].astype(float).sem()
            data.loc[(data['ID']==p), 'SC_E2_SE']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)'].astype(float).sem()
            data.loc[(data['ID']==p), 'SC_P4_SE']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)'].astype(float).sem()
            data['SF_E2:P4_Ratio_mean']=data['SF_E2_mean']/data['SF_P4_mean']
            data['SC_E2:P4_Ratio_mean']=data['SC_E2_mean']/data['SC_P4_mean']
            data['SF_E2:P4_Ratio_std']=data['SF_E2_std']/data['SF_P4_std']
            data['SC_E2:P4_Ratio_std']=data['SC_E2_std']/data['SC_P4_std']
            
            data.loc[(data['ID']==p), 'SF_E2_linreg_y_pred_max']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)_y_pred'].max()
            data.loc[(data['ID']==p), 'SC_E2_linreg_y_pred_max']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)_y_pred'].max()
            data.loc[(data['ID']==p), 'SF_P4_linreg_y_pred_max']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)_y_pred'].max()
            data.loc[(data['ID']==p), 'SC_P4_linreg_y_pred_max']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)_y_pred'].max()
            data.loc[(data['ID']==p), 'SF_E2_linreg_y_pred_mean']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)_y_pred'].mean()
            data.loc[(data['ID']==p), 'SC_E2_linreg_y_pred_mean']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)_y_pred'].mean()
            data.loc[(data['ID']==p), 'SF_P4_linreg_y_pred_mean']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)_y_pred'].mean()
            data.loc[(data['ID']==p), 'SC_P4_linreg_y_pred_mean']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)_y_pred'].mean()
            data.loc[(data['ID']==p), 'SF_E2_linreg_y_pred_sd']=data2.loc[(data2['ID']==p), 'SF_Estradiol (pg/mL)_y_pred'].std()
            data.loc[(data['ID']==p), 'SC_E2_linreg_y_pred_sd']=data2.loc[(data2['ID']==p), 'SC_Estradiol (pg/mL)_y_pred'].std()
            data.loc[(data['ID']==p), 'SF_P4_linreg_y_pred_sd']=data2.loc[(data2['ID']==p), 'SF_Progesterone (pg/mL)_y_pred'].std()
            data.loc[(data['ID']==p), 'SC_P4_linreg_y_pred_sd']=data2.loc[(data2['ID']==p), 'SC_Progesterone (pg/mL)_y_pred'].std()
            
            data['Subgroup#']=data['Subgroup'].map({'LOW RESPONSE':1, 'AVERAGE RESPONSE':2, 'HIGH RESPONSE':3})


            ##BC TRIGGER ==1 copy column twice
            

        
        #data.to_csv(path+'machine_learning_dataframe.csv')
        return data

    def synthetics(self, data):
        print(data['BC_P4_Decision'].value_counts())
        datapos = data.loc[data['BC_P4_Decision']==2][:12].copy()
        datapos2 = data.loc[data['BC_P4_Decision']==1][:6].copy()
        datapos3 = data.loc[data['BC_P4_Decision']==1][:2].copy()
        datapos2 = datapos2.append(datapos3.copy())

        data2 = pd.concat([data, datapos, datapos2], ignore_index=True)
        return data2

    def fit_rf_grid(self, data):  
        path = self.path



        #print('RF grid search path is: ', path)

        ml_dict = {'Classifier':[], 'Hormone target':[], 'X variables': [], 'Accuracy':[], 'Precision':[], 'Recall':[], 'Specificity':[], 'Crossvalidation Accuracy':[], 
        'TP':[], 'TN':[], 'FP':[], 'FN':[]}

        print(data['BC_P4_Decision'].value_counts())
        
        data.fillna(0, inplace=True)

        #data['BC_E2_Decision'].replace(to_replace=np.nan, value=0, inplace=True)
        data['BC_P4_Decision'].replace(to_replace=np.nan, value=0, inplace=True)
        # if column age exists, convert to int

        
        
        if 'BC_E2_Decision' in data.columns:
            data['BC_E2_Decision'].replace(to_replace=np.nan, value=0, inplace=True)
            data['BC_E2_Decision'] = data['BC_E2_Decision'].astype('int')

            xdata=data.drop(['BC_E2_Decision', 'BC_P4_Decision', 'Subgroup'], axis=1)

        else:
            xdata=data.drop(['BC_P4_Decision', 'Subgroup'], axis=1)

        
        #'SF_E2_linreg_y_pred_r_max', 'SC_E2_linreg_y_pred_r_max', 'SF_P4_linreg_y_pred_r_max', 'SC_P4_linreg_y_pred_r_max'], axis=1) 'BC_P4_Decision_1000', 'BC_E2_Decision_1000', 'BC_P4_Decision_f', 'BC_E2_Decision_f', 

        #KEEP ALL COLUMNS BEGINNING WITH SC
        sc_col = [x for x in xdata.columns if (x.startswith('SC')) | (x.startswith('Saliva'))| (x == 'Subgroup#') | (x.startswith('BC_E2_S1')) | (x.startswith('BC_P4_S1')) | (x=='Stimulation') | (x=='Age') | (x=='BMI')]
        sf_col = [x for x in xdata.columns if (x.startswith('SF')) | (x.startswith('Saliva'))| (x == 'Subgroup#') | (x.startswith('BC_E2_S1')) | (x.startswith('BC_P4_S1')) | (x=='Stimulation') | (x=='Age') | (x=='BMI')]
        all_col = [x for x in xdata.columns if x.startswith('S')]

        ys =['BC_P4_Decision']
        #features = xdata.columns


        x_opts=[sf_col, sc_col]

        svc=SVC(kernel='linear', C=0.1)
        logreglinear=LogisticRegression(solver='liblinear', C=0.1)
        logreg=LogisticRegression(solver='lbfgs', C=0.1)
        knn=KNN(n_neighbors=3)
        trees=tree.DecisionTreeClassifier()
        rf=RF()
        bagging = BGC(KNN(),max_samples=0.5, max_features=0.5) 
        gbc=GBC(n_estimators=10)
        abc=ABC(n_estimators=10)
        pca = decomposition.PCA()
        classifiers = [svc, logreglinear, logreg, knn, trees, rf, bagging, gbc, abc]
        doc = Document()
        doc.add_heading('Classifier: ' + str(rf), 0)
        print(data['BC_P4_Decision'].value_counts())
        for xo in x_opts:
            for target in ys:
                
                
                x= data[xo].fillna(data[xo].median())


                #x = preprocessing.scale(x)
                #preprocessing scaler so each variable is comparable later on
                y=data[target]
                #print(y)
                #n_components = list(range(1,x.shape[1]+1,1))
                criterion = ['gini', 'entropy']
                max_depth = [3,4,6,7,8,9,10,12]


                X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.44, random_state=42)

                

                

                # Number of trees in random forest
                n_estimators = [int(x) for x in np.linspace(start = 400, stop = 1500, num = 10)]
                # Number of features to consider at every split
                max_features = ['auto', 'sqrt']
                # Maximum number of levels in tree
                max_depth = [int(x) for x in np.linspace(1, 30, num = 11)]
                max_depth.append(None)
                # Minimum number of samples required to split a node
                min_samples_split = [2, 5, 10]
                # Minimum number of samples required at each leaf node
                min_samples_leaf = [1, 2, 4]
                # Method of selecting samples for training each tree
                bootstrap = [True, False]
                # Create the random grid
                random_grid = {'n_estimators': n_estimators,
                            'max_features': max_features,
                            'max_depth': max_depth,
                            'min_samples_split': min_samples_split,
                            'min_samples_leaf': min_samples_leaf,
                            'bootstrap': bootstrap}
            
                

                rf_random = RandomizedSearchCV(estimator = rf, param_distributions = random_grid, n_iter = 100, random_state=42, n_jobs = -1, 
                                            scoring = ('f1_micro'), refit='f1_micro', cv = 3, verbose=2, return_train_score=True)

                #print(data['BC_P4_Decision'].value_counts())
                
                rf_random.fit(X_train, y_train)
                print(target.split('_')[1], xo[1][:2])
                clf=rf_random
                i = clf.best_index_
                '''print('best index:', i)
                best_precision = clf.cv_results_['mean_test_precision'][i]
                best_recall = clf.cv_results_['mean_test_recall'][i]
                
                
                print('Best score (accuracy): {}'.format(clf.best_score_))
                print('Mean precision: {}'.format(best_precision))
                print('Mean recall: {}'.format(best_recall))
                print('Best parameters: {}'.format(clf.best_params_))
                print('Best estimator: {}'.format(clf.best_estimator_))
                print(rf_random.best_params_)'''
                classifier = rf_random.best_estimator_
                joblib.dump(classifier, path+"/rf_{}_{}.joblib".format(target.split('_')[1], xo[1][:2]))


    def test_rf_model(self, data):
        path = self.path
        
        data.fillna(0, inplace=True)

        #data['BC_E2_Decision_f'] = data['BC_E2_Decision']
        #data['BC_P4_Decision_f'] = data['BC_P4_Decision']

        #data['BC_E2_Decision'].replace(to_replace=np.nan, value=0, inplace=True)
        data['BC_P4_Decision'].replace(to_replace=np.nan, value=0, inplace=True)

        if 'BC_E2_Decision' in data.columns:
            data['BC_E2_Decision'].replace(to_replace=np.nan, value=0, inplace=True)
            data['BC_E2_Decision'] = data['BC_E2_Decision'].astype('int')

            xdata=data.drop(['BC_E2_Decision', 'BC_P4_Decision', 'Subgroup'], axis=1)

        else:
            xdata=data.drop(['BC_P4_Decision', 'Subgroup'], axis=1)
        #'SF_E2_linreg_y_pred_r_max', 'SC_E2_linreg_y_pred_r_max', 'SF_P4_linreg_y_pred_r_max', 'SC_P4_linreg_y_pred_r_max'], axis=1)  'BC_P4_Decision_1000', 'BC_E2_Decision_1000', 'BC_P4_Decision_f', 'BC_E2_Decision_f', 

        #KEEP ALL COLUMNS BEGINNING WITH SC
        sc_col = [x for x in xdata.columns if (x.startswith('SC')) | (x.startswith('Saliva'))| (x == 'Subgroup#') | (x.startswith('BC_E2_S1')) | (x.startswith('BC_P4_S1')) | (x=='Stimulation') | (x=='Age') | (x=='BMI')]
        sf_col = [x for x in xdata.columns if (x.startswith('SF')) | (x.startswith('Saliva'))| (x == 'Subgroup#') | (x.startswith('BC_E2_S1')) | (x.startswith('BC_P4_S1')) | (x=='Stimulation') | (x=='Age') | (x=='BMI')]
        all_col = [x for x in xdata.columns if x.startswith('S')]

        ys =['BC_P4_Decision']
        #features = xdata.columns

        x_opts=[sf_col]
        y_pred_dict = {'y_pred':[], 'y_test':[]}
        doc= Document()
        doc.add_heading('Random Forest', 0)
        rf_dict = {'target': [], 'X variables':[],'cross_val_#':[], 'accuracy': [], 'recall': [], 'specificity': [], 'cross_val_accuracy_mean': [], 
                
                'cross_val_specificity_mean': [], 'cross_val_sensitivity_mean': [], 'n': [], 
                'hyperparams': [], 'top10_feature_importances': []}
        for xo in x_opts:
            
            #for c in classifiers:

            for target in ys:
                x= data[xo].fillna(data[xo].median())


                #x = preprocessing.scale(x)
                #preprocessing scaler so each variable is comparable later on
                y=data[target]
                

                X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.4, random_state=100)
                classifier = joblib.load(path+"/rf_{}_{}.joblib".format(target.split('_')[1], xo[1][:2]))
                #print(classifier)

                y_pred = classifier.predict(X_test)
                ##y_pred AND x as column to dataframe
                y_pred_dict['y_pred'].append(list(y_pred))
                y_pred_dict['y_test'].append(list(y_test))
                #y_pred_dict['x_test'].append(list(X_test[0]))
                
                
                #print(y_pred)
                print(target.split('_')[1], xo[1][:2])
                '''cnf_matrix = metrics.confusion_matrix(y_test, y_pred)
                
                cnf = confusion_matrix(y_test, y_pred).ravel()
                #if len(cnf) == 4:
                tn, fp = cnf
                #print(cnf_matrix)
                specificity = tn / (tn + fp)
                #print('Specificity: ', specificity)
                #visualise confusion matrix
                class_names=[1, 2, 3] # name  of classes
                fig, ax = plt.subplots()
                tick_marks = np.arange(len(class_names))
                plt.xticks(tick_marks, class_names)
                plt.yticks(tick_marks, class_names)
                # create heatmap
                sns.heatmap(pd.DataFrame(cnf_matrix), annot=True, cmap='Spectral', fmt='g')
                ax.xaxis.set_label_position("top")
                plt.tight_layout()
                plt.title('Confusion matrix', y=1.1)
                plt.ylabel('Actual label')
                plt.xlabel('Predicted label')
                plt.tight_layout()
                plt.savefig(path+'confusion_matrix_{}_{}.png'.format(target.split('_')[1], xo[1][:2]), dpi=300, bbox_inches='tight', facecolor='white', transparent=False)
                plt.close()

                def confusion_matrix_scorer(clf, X, y):
                    y_pred = clf.predict(X)
                    cm = confusion_matrix(y, y_pred)
                    return {'tn': cm[0, 0], 'fp': cm[0, 1],
                            'fn': cm[1, 0], 'tp': cm[1, 1]}
                cv_results = cross_validate(classifier, x, y, cv=5,
                                            scoring=confusion_matrix_scorer)'''


                spec = make_scorer(recall_score, pos_label=0, average='macro')
                sensitivity = make_scorer(recall_score, pos_label=2, average='macro')
                scores=cross_val_score(classifier, x, y, cv=5, scoring='accuracy')

                print("Cross validation accuracy:", scores.mean())

                sp_scores=cross_val_score(classifier, x, y, cv=5, scoring=spec)

                print('specificity cross val:', sp_scores)

                print('specificity mean:', sp_scores.mean())
                print('specificity std:', sp_scores.std())

                se_scores=cross_val_score(classifier, x, y, cv=5, scoring=sensitivity)

                print('sensitivity cross val:', se_scores)
                print('sensitivity mean:', se_scores.mean())
                print('sensitivity std:', se_scores.std())
                #precision_scores=cross_val_score(classifier, x, y, cv=5, scoring='precision', average='weighted')
                #print('precision cross val:', precision_scores)
                #print('precision mean:', precision_scores.mean())
                #print('precision std:', precision_scores.std())

                from sklearn.metrics import precision_recall_fscore_support
                from sklearn.metrics import roc_auc_score

                print(target.split('_')[1], xo[1][:2])

                print(classification_report(y_test, y_pred))

                #roc_auc_score(y_test, y_pred, average="weighted", multi_class="ovr")

                #print('roc_auc_score:', roc_auc_score(y_test, y_pred, average="weighted", multi_class="ovr"))

                res = []
                for l in [0,1,2]:
                    prec,recall,_,_ = precision_recall_fscore_support(np.array(y_test)==l,
                                                      np.array(y_pred)==l,
                                                      pos_label=True,average=None)
                    res.append([l,recall[0],recall[1]])

                resdf = pd.DataFrame(res,columns = ['class','sensitivity','specificity'])
                print(resdf)

                '''res = []
                for l in [0,1,2]:
                    prec,recall,_,_ = precision_recall_fscore_support(np.array(y_test)==l,
                                                                np.array(y_pred)==l,
                                                                pos_label=True,average=None)
                    res.append([l,recall[0],recall[1], prec[0]])

                resdf = pd.DataFrame(res,columns = ['class','sensitivity','specificity','precision'])

                resdf.to_csv(path+'{}_{}_resdf.csv'.format(target.split('_')[1], xo[1][:2]), index=False)'''


                importances = classifier.feature_importances_
                feature_imp = pd.Series(importances, index = xo).sort_values(ascending = False)
                #print(feature_imp.index[:10])
                std = np.std([tree.feature_importances_ for tree in classifier.estimators_], axis=0)
                #forest_importances = pd.Series(importances, index=xo)
                fig, ax = plt.subplots(figsize=(20, 10), dpi=600)
                #features = pd.Series(feature_imp.index[:10])
                feature_imp.plot.bar(yerr=std, ax=ax)
                ax.set_title('Feature importance: {} Decision with {} X variables'.format(target.split('_')[1], xo[1][:2]), fontsize=16)
                #ax.set_title("Feature importances using MDI")
                ax.set_ylabel('Importance', fontsize=16)
                ax.set_xlabel('Features', fontsize=16)
                plt.xticks(fontsize=16)
                #fig.tight_layout()
                plt.tight_layout()
                plt.savefig(path+'feature_importance_{}_{}.png'.format(target.split('_')[1], xo[1][:2]), dpi=1200, bbox_inches='tight', facecolor='white', transparent=False)
                plt.show()
                plt.close()


                #shapval = shap.TreeExplainer(classifier).shap_values(X_test)
                #exp = shap.LinearExplainer(logreg, x)
                #shapval=explainer.shap_values(X_test)
                #shap.summary_plot(shapval[0], X_test)
                #plot=shap.summary_plot(shapval, x)
                #plt.show()


                for i in range(0, 5):
                    rf_dict['target'].append(target.split('_')[1])
                    rf_dict['X variables'].append(xo[1][:2])
                    rf_dict['cross_val_#'].append(i)
                    rf_dict['accuracy'].append(list(scores)[i])
                    #rf_dict['precision'].append(list(precision_scores)[i]) 
                    rf_dict['recall'].append(list(se_scores)[i])
                    rf_dict['specificity'].append(sp_scores[i])
                    rf_dict['cross_val_accuracy_mean'].append(scores.mean())
                    rf_dict['cross_val_specificity_mean'].append(sp_scores.mean())
                    rf_dict['cross_val_sensitivity_mean'].append(se_scores.mean())

                    '''rf_dict['tp'].append(list(cv_results['test_tp'])[i]) 
                    rf_dict['tn'].append(list(cv_results['test_tn'])[i])
                    rf_dict['fp'].append(list(cv_results['test_fp'])[i])
                    rf_dict['fn'].append(list(cv_results['test_fn'])[i])'''
                    rf_dict['top10_feature_importances'].append(list(feature_imp.index[:10]))
                    rf_dict['n'].append(len(y))
                    rf_dict['hyperparams'].append(classifier.get_params())


                doc.add_heading('Predicting {} trigger decisions with X variables from {} measurements'.format(target.split('_')[1], xo[1][:2]), level=1)
                doc.add_paragraph('Accuracy: {}'.format(metrics.accuracy_score(y_test, y_pred)))
                doc.add_paragraph('Precision: {}'.format(metrics.precision_score(y_test, y_pred, zero_division=1, average='weighted')))
                doc.add_paragraph('Recall: {}'.format(metrics.recall_score(y_test, y_pred, zero_division=1, average='weighted')))
                doc.add_paragraph('Specificity: {}'.format(spec))
                doc.add_paragraph(' ')
                doc.add_paragraph('A 5-fold cross validation was implemented to assess the model performance. The mean scores are as follows:')
                doc.add_paragraph('Accuracy: {}'.format(scores.mean()))
                doc.add_paragraph('Specificity: {}'.format(sp_scores.mean()))
                doc.add_paragraph('Recall: {}'.format(se_scores.mean()))
                doc.add_paragraph(' ')
                doc.add_paragraph('The confusion matrix shows the number of true positives, true negatives, false positives and false negatives.')
                doc.add_picture(path+'confusion_matrix_{}_{}.png'.format(target.split('_')[1], xo[1][:2]), width=Inches(5))
                doc.add_paragraph(' ')
                doc.add_paragraph('The feature importance plot shows the relative importance of each feature in the model. The top 10 features are as follows:')
                doc.add_paragraph('{}'.format(list(feature_imp.index[:10])))
                doc.add_picture(path+'feature_importance_{}_{}.png'.format(target.split('_')[1], xo[1][:2]), width=Inches(5))
                doc.add_paragraph(' ')
                #doc.add_paragraph('A random forest algorithm is an ensemble classifier trained with a multiptude (often hundreds) of decision trees. The decision is the class selected by the most trees.')
                #doc.add_paragraph('The following plot shows an example of one of the decision trees used to build this model. It shows the decision path made by the model.')
                #doc.add_picture(path+'tree_{}_{}.png'.format(target.split('_')[1], xo[0][:2]), width=Inches(5))
                
            


                
                #print(rf_dict)
        rf_df = pd.DataFrame(rf_dict)
        rf_df.to_csv(path+'random_forest_results_3groups.csv')
        #print(rf_df)
        #path = 'C:/Users/KatherineRidley/Mint Diagnostics Ltd/Mint Drive - Documents/Data/Saliva_Blood_Comparison_Studies_KR/IVIRMA/Stats/'
        doc.save(path + 'random_forest_results_3groups.docx')
        convert(path+'random_forest_results_3groups.docx', path+'random_forest_results_3groups.pdf')

        return y_pred_dict

        


